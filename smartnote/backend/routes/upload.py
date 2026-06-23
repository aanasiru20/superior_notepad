from fastapi import APIRouter, UploadFile, File, HTTPException, status
from fastapi.responses import JSONResponse
import os
import shutil
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document
from models import FileExtraction, FileResponse

router = APIRouter()

# Create uploads directory if not exists
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# ============== PDF EXTRACTION ==============

def extract_text_from_pdf(file_path: str) -> dict:
    """Extract text from PDF file using pdfplumber"""
    try:
        extracted_text = ""
        page_count = 0
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    extracted_text += text + "\n"
        
        return {
            "text": extracted_text.strip(),
            "page_count": page_count,
            "success": True
        }
    except Exception as e:
        return {
            "text": "",
            "page_count": 0,
            "success": False,
            "error": str(e)
        }

def extract_text_from_pdf_fallback(file_path: str) -> dict:
    """Fallback: Extract text from PDF using PyPDF2"""
    try:
        extracted_text = ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text:
                    extracted_text += text + "\n"
        
        return {
            "text": extracted_text.strip(),
            "page_count": page_count,
            "success": True
        }
    except Exception as e:
        return {
            "text": "",
            "page_count": 0,
            "success": False,
            "error": str(e)
        }

# ============== DOCX EXTRACTION ==============

def extract_text_from_docx(file_path: str) -> dict:
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        extracted_text = ""
        
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_text += para.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        extracted_text += cell.text + " "
                extracted_text += "\n"
        
        return {
            "text": extracted_text.strip(),
            "page_count": len(doc.paragraphs),
            "success": True
        }
    except Exception as e:
        return {
            "text": "",
            "page_count": 0,
            "success": False,
            "error": str(e)
        }

# ============== FILE UPLOAD ENDPOINTS ==============

@router.post("/pdf")
async def upload_pdf(file: UploadFile = File(...)):
    """
    Upload and extract text from PDF file
    
    - **file**: PDF file (multipart/form-data)
    """
    try:
        # Validate file type
        if file.content_type not in ["application/pdf"]:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Invalid file type. Only PDF files are allowed."
            )
        
        # Save file temporarily
        file_path = UPLOAD_DIR / file.filename
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extract text
        result = extract_text_from_pdf(str(file_path))
        
        # Fallback if pdfplumber fails
        if not result["success"]:
            result = extract_text_from_pdf_fallback(str(file_path))
        
        if not result["success"]:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to extract text from PDF: {result.get('error', 'Unknown error')}"
            )
        
        return {
            "status": "success",
            "filename": file.filename,
            "extracted_text": result["text"],
            "page_count": result["page_count"],
            "file_size": file_path.stat().st_size,
            "file_path": str(file_path)
        }
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error uploading PDF: {str(e)}"
        )

@router.post("/docx")
async def upload_docx(file: UploadFile = File(...)):
    """
    Upload and extract text from DOCX file
    
    - **file**: DOCX file (multipart/form-data)
    """
    try:
        # Validate file type
        if file.content_type not in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Invalid file type. Only DOCX files are allowed."
            )
        
        # Save file temporarily
        file_path = UPLOAD_DIR / file.filename
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extract text
        result = extract_text_from_docx(str(file_path))
        
        if not result["success"]:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to extract text from DOCX: {result.get('error', 'Unknown error')}"
            )
        
        return {
            "status": "success",
            "filename": file.filename,
            "extracted_text": result["text"],
            "page_count": result["page_count"],
            "file_size": file_path.stat().st_size,
            "file_path": str(file_path)
        }
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error uploading DOCX: {str(e)}"
        )

@router.post("/any")
async def upload_any_file(file: UploadFile = File(...)):
    """
    Upload any file (PDF or DOCX) and auto-detect format
    
    - **file**: PDF or DOCX file (multipart/form-data)
    """
    try:
        # Determine file type by extension
        file_extension = Path(file.filename).suffix.lower()
        
        if file_extension == ".pdf":
            return await upload_pdf(file)
        elif file_extension == ".docx":
            return await upload_docx(file)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type: {file_extension}. Only PDF and DOCX are supported."
            )
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error uploading file: {str(e)}"
        )

# ============== FILE CLEANUP ==============

@router.delete("/cleanup")
async def cleanup_uploads():
    """
    Clean up temporary upload files (use with caution)
    """
    try:
        if UPLOAD_DIR.exists():
            shutil.rmtree(UPLOAD_DIR)
            UPLOAD_DIR.mkdir(exist_ok=True)
            return {
                "status": "success",
                "message": "Upload directory cleaned"
            }
        return {
            "status": "success",
            "message": "No upload directory to clean"
        }
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error cleaning up uploads: {str(e)}"
        )
